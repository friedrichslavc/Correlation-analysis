# 将Markdown文档转换为Word文档

import os
import sys
import subprocess

def install_required_packages():
    """安装必要的包"""
    packages = ['python-docx', 'markdown']
    for package in packages:
        try:
            # 尝试导入包，如果失败则安装
            __import__(package.replace('-', '_'))
            print(f"{package} 已安装")
        except ImportError:
            print(f"正在安装 {package}...")
            subprocess.check_call([sys.executable, "-m", "pip", "install", package])
            print(f"{package} 安装完成")

def convert_md_to_docx(md_file_path, docx_file_path=None):
    """将Markdown文件转换为Word文档"""
    if docx_file_path is None:
        # 如果没有指定输出路径，则使用与输入文件相同的名称但扩展名为.docx
        docx_file_path = os.path.splitext(md_file_path)[0] + '.docx'
    
    try:
        # 导入必要的库
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import markdown
        from bs4 import BeautifulSoup
        import re
        
        # 安装BeautifulSoup如果尚未安装
        try:
            __import__('bs4')
        except ImportError:
            subprocess.check_call([sys.executable, "-m", "pip", "install", "beautifulsoup4"])
            from bs4 import BeautifulSoup
        
        # 读取Markdown文件内容
        with open(md_file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # 将Markdown转换为HTML
        html_content = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])
        
        # 使用BeautifulSoup解析HTML
        # 首先尝试使用html.parser解析器
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # 确保HTML有一个有效的结构
        if not soup.body:
            # 如果没有body标签，创建一个完整的HTML结构
            html_content = f"<html><body>{html_content}</body></html>"
            soup = BeautifulSoup(html_content, 'html.parser')
        
        # 创建一个新的Word文档
        doc = Document()
        
        # 设置文档样式
        style = doc.styles['Normal']
        style.font.name = 'SimSun'  # 宋体
        style.font.size = Pt(12)
        
        # 检查soup.body是否存在
        if soup.body is None:
            # 如果body不存在，直接处理整个soup
            elements = soup.children
        else:
            # 如果body存在，处理body的子元素
            elements = soup.body.children
            
        # 按顺序处理所有元素
        for element in elements:
            # 跳过纯文本节点
            if element.name is None:
                continue
                
            # 处理标题
            if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                heading_level = int(element.name[1])
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(element.get_text())
                run.bold = True
                run.font.size = Pt(18 - heading_level * 2)  # h1=16pt, h2=14pt, ...
            
            # 处理段落
            elif element.name == 'p':
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.add_run(element.get_text())
            
            # 处理无序列表
            elif element.name == 'ul':
                for li in element.find_all('li', recursive=False):
                    p = doc.add_paragraph()
                    p.style = 'List Bullet'
                    p.add_run(li.get_text())
            
            # 处理有序列表
            elif element.name == 'ol':
                for i, li in enumerate(element.find_all('li', recursive=False), 1):
                    p = doc.add_paragraph()
                    p.style = 'List Number'
                    p.add_run(li.get_text())
            
            # 处理表格
            elif element.name == 'table':
                rows = element.find_all('tr')
                if rows:
                    # 获取表头
                    headers = [th.get_text() for th in rows[0].find_all(['th', 'td'])]
                    
                    # 创建表格
                    table_rows = len(rows)
                    table_cols = len(headers)
                    doc_table = doc.add_table(rows=table_rows, cols=table_cols)
                    doc_table.style = 'Table Grid'
                    
                    # 填充表头
                    for i, header in enumerate(headers):
                        doc_table.cell(0, i).text = header
                    
                    # 填充表格内容
                    for i in range(1, table_rows):
                        cells = rows[i].find_all(['td', 'th'])
                        for j, cell in enumerate(cells):
                            if j < table_cols:  # 确保不超出列数
                                doc_table.cell(i, j).text = cell.get_text()
            
            # 处理代码块
            elif element.name == 'pre':
                code = element.find('code')
                if code:
                    p = doc.add_paragraph()
                    p.style = 'No Spacing'
                    code_text = code.get_text()
                    run = p.add_run(code_text)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
        
        # 处理图片
        for img in soup.find_all('img'):
            src = img.get('src')
            if src and os.path.exists(src):
                doc.add_picture(src, width=Inches(6.0))
        
        # 保存文档
        doc.save(docx_file_path)
        print(f"转换完成！文档已保存为: {docx_file_path}")
        return docx_file_path
    
    except Exception as e:
        print(f"转换过程中出错: {str(e)}")
        return None

def main():
    # 安装必要的包
    install_required_packages()
    
    # 获取当前目录
    current_dir = os.path.dirname(os.path.abspath(__file__))
    
    # 设置Markdown文件路径
    md_file_path = os.path.join(current_dir, "关联规则分析报告.md")
    
    # 设置Word文档输出路径
    docx_file_path = os.path.join(current_dir, "关联规则分析报告.docx")
    
    # 转换文件
    convert_md_to_docx(md_file_path, docx_file_path)

if __name__ == "__main__":
    main()